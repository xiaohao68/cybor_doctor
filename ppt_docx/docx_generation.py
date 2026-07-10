'''将大模型生成的json数据转换为word，可修改代码自定义word的样式'''
import os
import re
import hashlib
import time

from typing import Dict
from env import get_app_root

_OUTPUT_DIR_DOCX = os.path.join(get_app_root(), "data/cache/docx")

# 如果文件夹路径不存在，先创建
if not os.path.exists(_OUTPUT_DIR_DOCX):
    os.makedirs(_OUTPUT_DIR_DOCX)

def get_file_path_docx(text):
    """生成唯一的文件路径"""
    file_name = hashlib.sha256(text.encode("utf-8")).hexdigest()  # 可以使用uuid替代
    return os.path.join(_OUTPUT_DIR_DOCX, f"{file_name}.docx")

def is_chinese(text: str) -> bool:
    """判断文本中是否包含中文字符"""
    return bool(re.search(r'[\u4e00-\u9fff]', text))


from docx.shared import Pt
#qn函数用于创建XML命名空间的QName对象，解决字体设置问题，确保在不同区域设置字体时，字体能够正常显示
from docx.oxml.ns import qn 

def generate_docx_content(docx_content: Dict) -> str:
    """生成 docx 文件"""
    #Document类用于创建docx文档对象
    from docx import Document
    document = Document()

    # Word 标题 add_heading方法用于添加标题，参数为标题文本和标题级别
    title_heading = document.add_heading(docx_content['title'], 0)

    #alignment属性用于设置段落对齐方式  WD_PARAGRAPH_ALIGNMENT.CENTER 为居中对齐
    #WD_PARAGRAPH_ALIGNMENT枚举类用于设置段落对齐方式
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #runs属性用于获取段落中的所有运行对象
    title_run = title_heading.runs[0]
    
    # 根据标题是否包含中文设置字体
    if is_chinese(docx_content['title']):
        #font.name属性用于设置字体名称
        #_element属性用于获取运行对象的XML元素 rPr元素用于设置段落格式 rFonts元素用于设置字体 
        #qn函数用于创建XML命名空间的QName对象 w:eastAsia 表示东亚区域的字体设置
        title_run.font.name = '黑体'  # 中文字体
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 设置中文字体
    else:
        #font表示运行对象的字体属性
        title_run.font.name = 'Arial'  
    
    title_run.font.size = Pt(24)  # 标题字体大小

    # 页内容
    print(f'总共 {len(docx_content["sections"])} 个章节')
    for i, section in enumerate(docx_content['sections']):
        print(f'生成第 {i + 1} 章节: {section["heading"]}')
        section_heading = document.add_heading(section['heading'], level=1)
        #runs属性用于获取段落中的所有运行对象
        section_heading_run = section_heading.runs[0]
        
        # 根据章节标题是否包含中文设置字体
        if is_chinese(section['heading']):
            section_heading_run.font.name = '宋体'
            section_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            section_heading_run.font.name = 'Times New Roman'
        
        section_heading_run.font.size = Pt(16)  # 章节标题字体大小

        for paragraph in section['paragraphs']:
            para_heading = document.add_heading(paragraph['heading'], level=2)
            para_heading_run = para_heading.runs[0]
            
            # 根据段落标题是否包含中文设置字体
            if is_chinese(paragraph['heading']):
                para_heading_run.font.name = '宋体'
                para_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                para_heading_run.font.name = 'Calibri'
            
            para_heading_run.font.size = Pt(14)  # 段落标题字体大小

            # 设置正文内容字体 add_paragraph方法用于添加段落，参数为段落文本
            p = document.add_paragraph(paragraph['content'])
            #runs属性用于获取段落中的所有运行对象
            p_run = p.runs[0]
            
            # 根据正文内容是否包含中文设置字体
            if is_chinese(paragraph['content']):
                p_run.font.name = '宋体'
                p_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                p_run.font.name = 'Arial'
            
            p_run.font.size = Pt(12)  # 正文字体大小

    _output_file = get_file_path_docx(str(time.time()))
    document.save(_output_file)

    return _output_file

